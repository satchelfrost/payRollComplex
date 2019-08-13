import openpyxl as opx
import calendar
from docx import Document
import datetime
import re

'''Converts month number to month name'''
def Month(num):
    months = {
        1 : "January", 
     	2 : "February", 
    	3 : "March", 
        4 : "April", 
     	5 : "May", 
     	6 : "June", 
    	7 : "July",
        8 : "August",
     	9 : "September", 
    	10 : "October", 
        11 : "November",
    	12 : "December" }
    return months[num]

'''Uses regular expressions to parse dates into datetime'''
def parseDate(date):
    #re is the imported module regular expressions
    dateRegex = re.compile(r'(\d{2}) (\d{2}) (\d{4})')
    mo = dateRegex.search(date)
    month = mo.group(1)
    day = mo.group(2)
    year = mo.group(3)
    return datetime.datetime(int(year),int(month),int(day),0,0)

class payRoll:
    
    '''Loads employee sheet, finds date index, makes date list'''
    def __init__(self,file,employee,startDate,endDate):
        wb = opx.load_workbook(file+'.xlsx',data_only=True)
        self.employeeName = employee
        #loads employee sheet
        self.employee = wb[employee]
        #finds date index
        for i in range(71,131):
            date = self.employee.cell(row=i,column=2).value
            if date == startDate: self.dateDex = i
        #generates dates from start to end
        self.dates = []
        while(startDate <= endDate):
            self.dates.append(startDate)
            #effectively date++
            startDate += datetime.timedelta(days=1)
            
    '''Converts datetime dates to readable format'''
    def dateConvert(self):
        self.conv_dates = []
        for i in range(len(self.dates)):
            datename = calendar.day_name[self.dates[i].weekday()]
            day = str(self.dates[i].day)
            month = Month(self.dates[i].month)
            combined = datename + ', ' + month + ', ' + day
            self.conv_dates.append(combined)
            
    '''Loads the hours worked for both weeks'''
    def loadHours(self):
        self.hours = []
        #load first weeks hours worked
        for i in range(4,11):
            hour = self.employee.cell(row=self.dateDex,column=i).value
            self.hours.append(hour)
        #load second weeks hours worked
        for i in range(4,11):
            hour = self.employee.cell(row=self.dateDex+2,column=i).value
            self.hours.append(hour)    

    '''Loads the tips for both weeks'''
    def loadTips(self):
        self.tips = []
        #load first weeks tips
        for i in range(4,11):
            tip = self.employee.cell(row=self.dateDex+1,column=i).value
            self.tips.append(tip)
        #load second weeks tips
        for i in range(4,11):
            tip = self.employee.cell(row=self.dateDex+3,column=i).value
            self.tips.append(tip)
            
    '''Loads and calculates total pay earned'''
    def loadTotal(self):
        t1 = self.employee.cell(row=self.dateDex,column=14).value
        t2 = self.employee.cell(row=self.dateDex+2,column=14).value
        self.total = round(t1+t2,2)
        
    '''Removes zeros from list when a day was not worked'''   
    def cutZeros(self):
        self.hours_corr = []
        self.tips_corr = []
        self.dates_corr = []
        for i in range(len(self.hours)):
            #assumes if 0 hours worked then 0 tips as well
            if self.hours[i] != 0:
                self.hours_corr.append(self.hours[i])
                #add tips corrected with two dec. places
                self.tips_corr.append('{0:.2f}'.format(self.tips[i]))
                #note we are using converted dates i.e. readable
                self.dates_corr.append(self.conv_dates[i])
                
    '''Creates the actual word document'''         
    def makeDoc(self):
        self.dateConvert()
        self.loadHours()
        self.loadTips()
        self.loadTotal()
        self.cutZeros()
        document = Document()
        text = 'Hours worked for ' + self.employeeName
        document.add_heading(text, 0)
        #table heading
        table = document.add_table(rows=1, cols=3, style='Medium Shading 1')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Hours'
        hdr_cells[2].text = 'Tips'
        #adds table cells
        for i in range(len(self.hours_corr)):
            row_cells = table.add_row().cells
            row_cells[0].text = self.dates_corr[i]
            row_cells[1].text = str(self.hours_corr[i])
            row_cells[2].text = str(self.tips_corr[i])
        #table without borders to hold total pay
        table2 = document.add_table(rows=1, cols=3)
        row_cells = table2.add_row().cells
        row_cells[2].text = 'Total Pay: ' + str(self.total)
        signature = '\nSignature ________________________    Date ________________'
        document.add_paragraph(signature)
        document.save('payroll.docx')
        

def main():
    file = input('What is the name of the file? ')
    employee = input('Which employee? ')
    startDate = input('What is the start date? mm dd yyyy ')
    endDate = input('What is the end date? mm dd yyyy ')
    startDate = parseDate(startDate)
    endDate = parseDate(endDate)

    PR = payRoll(file,employee,startDate,endDate)
    PR.makeDoc()
    

if __name__ == '__main__': main()
